from pathlib import Path

# Generate DOCX
from docx import Document
# Generate PDF
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
# Generate image
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).parent
INP = ROOT / "inputs"
INP.mkdir(parents=True, exist_ok=True)

# Text sample
(INP / "text_idea.txt").write_text(
    """
Design a smartwatch companion app to manage health metrics and notifications.
Must support step tracking, heart-rate monitoring, and sleep analysis.
Should integrate with existing wearable APIs.
Deliver weekly summary reports and export data as CSV.
Constraints: Android 11+, iOS 15+, limited budget; initial prototype in 4 weeks.
""".strip(),
    encoding="utf-8",
)

# DOCX sample
specs = Document()
specs.add_heading("Product Specifications", level=1)
specs.add_paragraph("Purpose: Build a cross-platform mobile app for personal budgeting.")
specs.add_paragraph("Requirements:")
specs.add_paragraph("- Must import bank statements and categorize expenses.")
specs.add_paragraph("- Should provide savings goals and alerts.")
specs.add_paragraph("Constraints: Only React Native, at most 2GB data.")
specs.add_paragraph("Deliverables: Prototype, API spec, and user guide.")
specs.save(str(INP / "specs.docx"))

# PDF sample
pdf_path = INP / "requirements.pdf"
c = canvas.Canvas(str(pdf_path), pagesize=letter)
c.drawString(72, 720, "Build a web dashboard for IoT device telemetry.")
c.drawString(72, 700, "Must support real-time charts and alerting.")
c.drawString(72, 680, "Constraints: No external DB; file-based storage only.")
c.drawString(72, 660, "Deliverables: Setup scripts and deployment instructions.")
c.showPage()
c.save()

# Image sample (wireframe)
img = Image.new("RGB", (600, 400), color=(240, 240, 240))
draw = ImageDraw.Draw(img)
draw.rectangle([50, 50, 550, 350], outline=(0, 0, 0), width=2)
draw.text((60, 60), "Login Screen\nEmail\nPassword\nSign In", fill=(10, 10, 10))
img.save(str(INP / "wireframe.png"))

print(f"Samples generated in: {INP}")
